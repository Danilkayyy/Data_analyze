import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
from solution import Camera
import pprint
import datetime as dt
import docx
from pathlib import PurePosixPath, PureWindowsPath


class NotFoundFunction(Exception):
    pass


def process(path_to_file: str, detailed=1, frequency=30) -> None:
    data = np.load(path_to_file)

    right_leg = []
    left_leg = []

    [(right_leg.append((i[24][:2], i[26][:2], i[28][:2], i[30][:2], i[32][:2])),
      left_leg.append((i[23][:2], i[25][:2], i[27][:2], i[29][:2], i[31][:2]))) for i in data]

    right_leg = np.array(right_leg)
    left_leg = np.array(left_leg)
    right_leg = right_leg.reshape(right_leg.shape[0], 10)
    left_leg = left_leg.reshape(left_leg.shape[0], 10)

    right_leg_dataFrame = pd.DataFrame(data=right_leg,
                                       columns=["x_hip", "y_hip", "x_knee", "y_knee", "x_ankle", "y_ankle", "x_heel",
                                                "y_heel", "x_footIndex", "y_footIndex"])
    left_leg_dataFrame = pd.DataFrame(data=left_leg,
                                      columns=["x_hip", "y_hip", "x_knee", "y_knee", "x_ankle", "y_ankle", "x_heel",
                                               "y_heel", "x_footIndex", "y_footIndex"])

    keys_pare_XY = [right_leg_dataFrame.keys()[i:i + 2] for i in range(0, len(right_leg_dataFrame.keys()), 2)]
    for i in keys_pare_XY:
        print("\t\tx\t\t\ty")
        print(i[0].split("_")[1].upper())
        print("RIGHT")
        print("mean:", right_leg_dataFrame[i[0]].mean(), right_leg_dataFrame[i[1]].mean(), sep="\t")
        print("min:", right_leg_dataFrame[i[0]].min(), right_leg_dataFrame[i[1]].min(), sep="\t")
        print("max:", right_leg_dataFrame[i[0]].max(), right_leg_dataFrame[i[1]].max(), sep="\t")

        print("LEFT")
        print("mean:", left_leg_dataFrame[i[0]].mean(), left_leg_dataFrame[i[1]].mean(), sep="\t")
        print("min:", left_leg_dataFrame[i[0]].min(), left_leg_dataFrame[i[1]].min(), sep="\t")
        print("max:", left_leg_dataFrame[i[0]].max(), left_leg_dataFrame[i[1]].max(), sep="\t")
        print("######################################################################")

    if detailed not in range(1, 3):
        raise NotFoundFunction()

    if detailed == 1:
        get_all_graph(right_leg_dataFrame, left_leg_dataFrame, frequency)
    elif detailed == 2:
        get_union_graph(right_leg_dataFrame, left_leg_dataFrame, frequency)


def get_all_graph(right_leg_dataFrame: pd.DataFrame, left_leg_dataFrame: pd.DataFrame, frequency) -> None:
    fig, axs = plt.subplots(nrows=5, ncols=2)
    fig.suptitle("Detail process mediapipe data")

    keys_pare_XY = [right_leg_dataFrame.keys()[i:i + 2] for i in range(0, len(right_leg_dataFrame.keys()), 2)]

    for pareIndex, pare in enumerate(keys_pare_XY):
        axs[pareIndex][0].plot(np.arange(len(right_leg_dataFrame[pare[0]])) / frequency, right_leg_dataFrame[pare[0]],
                               label=f"right {pare[0]}", color="blue")
        axs[pareIndex][0].plot(np.arange(len(left_leg_dataFrame[pare[0]])) / frequency, left_leg_dataFrame[pare[0]],
                               label=f"left {pare[0]}", color="red")
        axs[pareIndex][0].legend()
        axs[pareIndex][0].grid()

        axs[pareIndex][1].plot(np.arange(len(right_leg_dataFrame[pare[1]])) / frequency, right_leg_dataFrame[pare[1]],
                               label=f"right {pare[1]}", color="blue")
        axs[pareIndex][1].plot(np.arange(len(left_leg_dataFrame[pare[1]])) / frequency, left_leg_dataFrame[pare[1]],
                               label=f"left {pare[1]}", color="red")
        axs[pareIndex][1].legend()
        axs[pareIndex][1].grid()
    plt.show()


def get_union_graph(right_leg_dataFrame: pd.DataFrame, left_leg_dataFrame: pd.DataFrame, frequency):
    fig, axs = plt.subplots(nrows=2, ncols=1)
    fig.suptitle("Union mediapipe data")

    axs[0].plot(np.arange(len(right_leg_dataFrame["x_hip"])) / frequency, right_leg_dataFrame["x_hip"],
                label="x_right_hip", color="blue")
    axs[0].plot(np.arange(len(left_leg_dataFrame["x_hip"])) / frequency, left_leg_dataFrame["x_hip"],
                label="x_left_hip", color="red")

    axs[0].plot(np.arange(len(right_leg_dataFrame["x_knee"])) / frequency, right_leg_dataFrame["x_knee"],
                label="x_right_knee", color="green")
    axs[0].plot(np.arange(len(left_leg_dataFrame["x_knee"])) / frequency, left_leg_dataFrame["x_knee"],
                label="x_left_knee", color="brown")

    axs[0].plot(np.arange(len(right_leg_dataFrame["x_ankle"])) / frequency, right_leg_dataFrame["x_ankle"],
                label="x_right_ankle", color="grey")
    axs[0].plot(np.arange(len(left_leg_dataFrame["x_ankle"])) / frequency, left_leg_dataFrame["x_ankle"],
                label="x_left_ankle", color="yellow")

    axs[0].legend()
    axs[0].grid()

    axs[1].plot(np.arange(len(right_leg_dataFrame["y_hip"])) / frequency, right_leg_dataFrame["y_hip"],
                label="y_right_hip", color="blue")
    axs[1].plot(np.arange(len(left_leg_dataFrame["y_hip"])) / frequency, left_leg_dataFrame["y_hip"],
                label="y_left_hip", color="red")

    axs[1].plot(np.arange(len(right_leg_dataFrame["y_knee"])) / frequency, right_leg_dataFrame["y_knee"],
                label="y_right_knee", color="green")
    axs[1].plot(np.arange(len(left_leg_dataFrame["y_knee"])) / frequency, left_leg_dataFrame["y_knee"],
                label="y_left_knee", color="brown")

    axs[1].plot(np.arange(len(right_leg_dataFrame["y_ankle"])) / frequency, right_leg_dataFrame["y_ankle"],
                label="y_right_ankle", color="grey")
    axs[1].plot(np.arange(len(left_leg_dataFrame["y_ankle"])) / frequency, left_leg_dataFrame["y_ankle"],
                label="y_left_ankle", color="yellow")

    axs[1].legend()
    axs[1].grid()

    plt.show()


def statistics(file_path: str) -> dict:
    """
    Получение среднего значения по координате Z, дисперсии от Z,
    Среднего значения скорости, дисперсии от скорости
    """
    df = pd.read_csv(file_path, sep=";")
    df["times"] = df['time'].diff()
    df['zs'] = abs(df['z'].diff())
    speed = (df['zs'][1:] / df['times'][1:]).where(df['times'] != 0, 0)
    return {"mean_z": df['z'].mean(), "disp_z": df['z'].var(), "mean_speed": speed.mean(), "disp_speed": speed.var()}


def process_trackers() -> dict:
    """
    Обработка данных с трекеров  (спина)
    """
    trackers_stat = {"new_road": [], "old_road": []}
    path = "data/trackers"  # путь до данных с трекеров

    if os.name != "posix":
        path = PureWindowsPath(path)

    for root, _, files in os.walk(path):
        for file in files:
            if not file.endswith('.csv'):
                continue
            file_path = os.path.join(root, file)

            if os.name != "posix":
                file_path = PurePosixPath(file_path)

            test_num = file_path.split("/")
            tracker_name = "@".join(test_num[4:])
            if tracker_name.split("@")[0] == "Tracker-LHR-1761CD18":
                continue
            trackers_stat[test_num[2]] = trackers_stat.get(test_num[2], []) + [
                (tracker_name, statistics(file_path))]

    return trackers_stat


# def process_trackers() -> dict:
#     """
#     Обработка данных с трекеров (спина, ноги)
#     """

#     hipTrackName = ""
#     rHeelTrackName = ""
#     lHeelTrackName = ""

#     check = {hipTrackName: "Hip", rHeelTrackName: "rHeel", lHeelTrackName: "lHeel"}

#     trackers_stat = {"new_road": [], "old_road": []}

#     path = "data/trackers"  # путь до данных с трекеров

#     if os.uname().sysname != "posix":
#         path = PureWindowsPath(path)

#     for root, _, files in os.walk(path):
#         for file in files:
#             if not file.endswith('.csv'):
#                 continue
#             file_path = os.path.join(root, file)

#             if os.uname().sysname != "posix":
#                 file_path = PurePosixPath(file_path)

#             test_num = file_path.split("/")
#             tracker_name = "@".join(test_num[4:])
#             if tracker_name.split("@")[0] == "Tracker-LHR-1761CD18":
#                 continue
#             trackers_stat[test_num[2]] = trackers_stat.get(test_num[2], []) + [
#                 (check[tracker_name], statistics(file_path))]

#     return trackers_stat



def get_distance_between_heel() -> dict:
    """
    Получение максимального расстояния между пятками (через камеры)
    """
    paths = []
    points_path = "data/points"
    result = {"new_road": [], "old_road": []}
    for root, _, files in os.walk(points_path):
        for file in files:
            file_path = os.path.join(root, file)

            if os.name!= "posix":
                file_path = PurePosixPath(file_path)

            paths.append(file_path)
    for path in paths:
        test_num = path.split("/")

        if os.name != "posix":
            data = np.load(PureWindowsPath(path))
        else:
            data = np.load(path)

        testDt = path.split("/")[3].rstrip(".npy")
        mean_distance = np.array([abs(x[29][0] - x[30][0]) for x in data]).max()
        result[test_num[2]] = result.get(test_num[2], []) + [(testDt, mean_distance)]
    return result


def process_videos() -> None:
    """
    Получение точек mediapipe из видео
    """
    videos_path = "data/videos"

    if os.name != "posix":
        videos_path = PureWindowsPath(videos_path)

    paths = []
    for root, _, files in os.walk("data/videos"):
        for file in files:
            file_path = os.path.join(root, file)

            if os.name != "posix":
                file_path = PurePosixPath(file_path)

            paths.append(file_path)

    for path in paths:
        if "_simple" not in path:
            continue

        test_num = path.split("/")

        if os.name != "posix":
            path = PureWindowsPath(path)

        camera = Camera(path)
        camera.load_points_from_videos()

        file_name = test_num[3].split("_")[0]

        npy_path = f"data/points/{test_num[2]}/{file_name}"

        if os.name != "posix":
            npy_path = PureWindowsPath(npy_path)

        camera.save_file(npy_path)


def get_hip() -> dict:
    """
    Получение точек со спины (видео)
    """

    paths = []
    result = {"new_road": [], "old_road": []}

    points_path = "data/points"

    if os.name != "posix":
        points_path = PureWindowsPath(points_path)

    for root, _, files in os.walk(points_path):
        for file in files:
            file_path = os.path.join(root, file)

            if os.name != "posix":
                file_path = PurePosixPath(file_path)

            paths.append(file_path)
    for path in paths:
        if path.endswith(".npy"):
            test_num = path.split("/")
            result[test_num[2]] = result.get(test_num[2], []) + [
                (test_num[3].rstrip(".npy"), np.load(path)[:, 24, :][:, 0].mean())]
    return result


def create_table() -> None:
    """
    Построение таблиц в docx
    """

    trackers = process_trackers()

    distBetweenHeel = get_distance_between_heel()
    hip = get_hip()
    for road in ["new_road", "old_road"]:
        trackers[road] = sorted(trackers[road],
                                key=lambda x: dt.datetime.strptime(x[0].split("@")[1].rstrip(".csv"),
                                                                   "%Y-%m-%d_%I-%M-%S"))
        distBetweenHeel[road] = sorted(distBetweenHeel[road],
                                       key=lambda x: dt.datetime.strptime(x[0], "%I:%M:%S.%f"))
        hip[road] = sorted(hip[road],
                           key=lambda x: dt.datetime.strptime(x[0], "%I:%M:%S.%f"))

    doc = docx.Document()

    doc.add_paragraph(f"\nTrackers\n")

    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    header = ["№", "Mean Z", "Disp Z", "Mean Speed", "Disp Speed"]

    for j in range(len(header)):
        table.cell(0, j).text = header[j]

    test_index = 1
    for road in trackers:
        for data in trackers[road]:
            table.cell(test_index, 0).text = str(test_index)
            table.cell(test_index, 1).text = f"{data[1]['mean_z']:.6f}"
            table.cell(test_index, 2).text = f"{data[1]['disp_z']:.6f}"
            table.cell(test_index, 3).text = f"{data[1]['mean_speed']:.6f}"
            table.cell(test_index, 4).text = f"{data[1]['disp_speed']:.6f}"
            test_index += 1

    doc.add_paragraph(f"\nVideos\n")

    table = doc.add_table(rows=9, cols=4)

    header = ["№", "right_hip_x.mean", "right_hip_x.var", "max_distance"]

    for j in range(len(header)):
        table.cell(0, j).text = header[j]

    paths = ["data/points/new_road", "data/points/old_road"]

    if os.name != "posix":
        paths = [PureWindowsPath(path) for  path in paths]

    for road_index, path in enumerate(paths):
        for i, file_name in enumerate(os.listdir(path), start=1):
            if not file_name.endswith('.npy'):
                continue
            file_path = os.path.join(path, file_name)
            if os.name == "posix":
                file_path = PurePosixPath(file_path)

            array = np.load(file_path)
            right_hip = array[:, 24, :]
            right_hip_x = right_hip[:, 0]
            
            max_distance = np.array([abs(x[29][0] - x[30][0]) for x in array]).max()

            table.cell(i + road_index * 4, 0).text = str(i + road_index * 4)
            table.cell(i + road_index * 4, 1).text = f"{right_hip_x.mean():.6f}"
            table.cell(i + road_index * 4, 2).text = f"{right_hip_x.var():.6f}"
            table.cell(i + road_index * 4, 3).text = f"{max_distance:.6f}"

    doc.save("table.docx")


if __name__ == "__main__":
    create_table()

    # try:
    #     process("data_for_process.npy", detailed=2, frequency=30)
    # except NotFoundFunction:
    #     print("[-] Функция не найдена, выберите: 1, 2.")
    # except Exception as exp:
    #     print("[-]", exp)
