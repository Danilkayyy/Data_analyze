import joblib
import pandas as pd
from sklearn.metrics import mean_absolute_error
import docx


def process_file(file_name: str) -> list:
    result = joblib.load(file_name)
    with pd.option_context('display.max_rows', None, 'display.max_columns', 25, 'display.width', 4000,
                           'display.float_format', str):
        dataForDocx = []
        for r in [10, 15]:
            rc = result[(result["Q"] == r) & (result["W"] == r)]
            modelsS = result[(result["Q"] == r) & (result["W"] == r)]["model_speed_name"].unique()
            modelsP = result[(result["Q"] == r) & (result["W"] == r)]["model_name"].unique()

            for mp in modelsP:
                for m in modelsS:
                    dataForDocx.append((mp[0].split(".")[0],
                                        mp[1],
                                        m[0].split(".")[0],
                                        m[1],
                                        round(mean_absolute_error(
                                            rc[(rc["model_speed_name"] == m) & (rc["model_name"] == mp)]["result"],
                                            rc[(rc["model_speed_name"] == m) & (rc["model_name"] == mp)][
                                                "speed_only_Q"]), 3),
                                        round(mean_absolute_error(
                                            rc[(rc["model_speed_name"] == m) & (rc["model_name"] == mp)][
                                                "result_future"],
                                            rc[(rc["model_speed_name"] == m) & (rc["model_name"] == mp)][
                                                "speed_only_W"]), 3))
                                       )

        return dataForDocx


def write_to_docx(all_data: list) -> None:
    doc = docx.Document()

    for data_index, data in enumerate(all_data):

        doc.add_paragraph(f"\n{data_index}\n")
        ROWS = len(data) + 1
        COLS = 6

        table = doc.add_table(rows=ROWS, cols=COLS)
        table.style = 'Table Grid'
        header = ["Test", "Mean ± SD", "p-value"]

        for j in range(len(header)):
            table.cell(0, j).text = header[j]
        
        row = 1
        name_index = 1
        value_index = 1
        for rowIndex, rowContent in enumerate(data):
            if row == name_index:
                table.cell(row, 0).merge(table.cell(row + 5, 0))

                table.cell(row, 0).text = str(rowContent[0])
                name_index += 6

            if row == value_index:
                table.cell(row, 1).merge(table.cell(row + 35, 1))
                table.cell(row, 3).merge(table.cell(row + 35, 3))

                table.cell(row, 1).text = str(rowContent[1])
                table.cell(row, 3).text = str(rowContent[3])
                value_index += 36

            table.cell(row, 2).text = str(rowContent[2])
            table.cell(row, 4).text = str(rowContent[4])
            table.cell(row, 5).text = str(rowContent[5])
            row += 1

    doc.save("table1.docx")


# #2
# joblib.load(res,"result_trackers_2.pd") #PD 2 столбцов
# joblib.load(res,"result_cv_2.pd")


if __name__ == "__main__":
    fileNames = ["result_trackers.pd", "result_cv.pd"]
    allData = [process_file(file_name=fileName) for fileName in fileNames]
    write_to_docx(all_data=allData)
